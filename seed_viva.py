"""
Script de poblado para Viva - Workflow Manager
Crea: formularios en nodos, trámites de ejemplo, workflow nuevo, docs Word/Excel
"""
import requests, json, os, sys, time
from datetime import datetime

BASE = "http://localhost/api"

# ── IDs existentes ────────────────────────────────────────────────────────────
COMPANY_ID   = "6a29bf47519cfc072fd2f2ab"
DEPT_ATENCION = "6a2ae51f2593e12cddbe3166"   # Atención al Cliente
DEPT_SOPORTE  = "6a2ae59d2593e12cddbe3167"   # Soporte Técnico
ROLE_RECEP    = "6a2ae5aa2593e12cddbe3168"   # Recepcionista
ROLE_TEC      = "6a2ae5b92593e12cddbe3169"   # Técnico de Soporte
USER_ADMIN    = "6a29bf1d519cfc072fd2f2aa"   # Admin Viva
USER_JUAN     = "6a2ae5da2593e12cddbe316a"   # Juan Pérez (Recepcionista)
USER_CARLOS   = "6a2ae5f22593e12cddbe316b"   # Carlos López (Técnico)
WF_INCIDENCIA = "6a2a2181519cfc072fd2f2ac"   # Workflow existente

# Nodos del workflow de incidencia
N_INICIO   = "6a2a218a519cfc072fd2f2ad"
N_RECEP    = "6a2a2191519cfc072fd2f2ae"
N_DIAG     = "6a2ae7bc2593e12cddbe316d"
N_DECISION = "6a2ae8622593e12cddbe3170"
N_CIERRE   = "6a2ae8db2593e12cddbe3172"
N_PROG     = "6a2aea4d0c92b72197936abc"

# Transiciones
T_INICIO_TO_RECEP   = "6a2a219e519cfc072fd2f2b0"
T_RECEP_TO_DIAG     = "6a2ae7c42593e12cddbe316e"
T_DIAG_TO_DEC       = "6a2ae86c2593e12cddbe3171"
T_DEC_ACEPTAR       = "6a2ae9d30c92b72197936ab8"
T_DEC_RECHAZAR      = "6a2aea520c92b72197936abd"
T_CIERRE_TO_FIN     = "6a2aea3f0c92b72197936abb"
T_PROG_TO_FIN       = "6a2aea9d0c92b72197936abf"

# ── Auth ──────────────────────────────────────────────────────────────────────
def login(email, password):
    r = requests.post(f"{BASE}/auth/login", json={"email": email, "password": password})
    r.raise_for_status()
    return r.json()["accessToken"]

def h(token):
    return {"Authorization": f"Bearer {token}"}

def post(url, token, body):
    r = requests.post(url, json=body, headers=h(token))
    if not r.ok:
        print(f"  ERROR {r.status_code}: {r.text[:200]}")
        return None
    return r.json()

def put(url, token, body):
    r = requests.put(url, json=body, headers=h(token))
    if not r.ok:
        print(f"  ERROR {r.status_code}: {r.text[:200]}")
        return None
    return r.json()

def upload_file(token, file_path, workflow_id):
    """Sube un archivo usando el endpoint legacy (workflowId). Devuelve metadata para formData."""
    with open(file_path, "rb") as f:
        ext = os.path.splitext(file_path)[1].lower()
        mime = {
            ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            ".pdf":  "application/pdf",
        }.get(ext, "application/octet-stream")
        files_data = {"file": (os.path.basename(file_path), f, mime)}
        data       = {"workflowId": workflow_id}
        r = requests.post(f"{BASE}/files/upload", files=files_data, data=data, headers=h(token))
    if not r.ok:
        print(f"  ERROR upload {r.status_code}: {r.text[:120]}")
        return None
    return r.json()

# ── Generar documentos de ejemplo ────────────────────────────────────────────
def gen_docs(out_dir="/tmp/viva_docs"):
    os.makedirs(out_dir, exist_ok=True)

    # Word - Reporte de Incidencia
    from docx import Document
    from docx.shared import Pt, RGBColor
    doc = Document()
    doc.add_heading("Reporte de Incidencia de Internet", 0)
    doc.add_paragraph(f"Fecha: {datetime.now().strftime('%d/%m/%Y')}")
    doc.add_heading("Datos del Cliente", 1)
    tabla = doc.add_table(rows=5, cols=2)
    tabla.style = "Table Grid"
    datos = [("Nombre","Juan García"),("Teléfono","71234567"),
             ("Dirección","Av. Arce 1234"),("Tipo de falla","Sin señal"),
             ("Plan","Fibra 100 Mbps")]
    for i,(k,v) in enumerate(datos):
        tabla.rows[i].cells[0].text = k
        tabla.rows[i].cells[1].text = v
    doc.add_heading("Descripción del Problema", 1)
    doc.add_paragraph("El cliente reporta pérdida total de conexión desde las 08:00 hrs. "
                      "Se verificó que el equipo enciende pero no conecta al servidor.")
    path_word = f"{out_dir}/reporte_incidencia.docx"
    doc.save(path_word)
    print(f"  Word: {path_word}")

    # Word - Contrato de Servicio
    doc2 = Document()
    doc2.add_heading("Contrato de Servicio de Internet", 0)
    doc2.add_paragraph("CONTRATO N° VS-2026-001")
    doc2.add_paragraph(f"Fecha de firma: {datetime.now().strftime('%d/%m/%Y')}")
    doc2.add_heading("Partes", 1)
    doc2.add_paragraph("PROVEEDOR: Viva S.A.\nCLIENTE: María López\nCI: 5678901")
    doc2.add_heading("Objeto del Contrato", 1)
    doc2.add_paragraph("Prestación del servicio de internet de banda ancha con velocidad "
                       "de 50 Mbps por un período de 12 meses.")
    doc2.add_heading("Cláusulas", 1)
    for i,c in enumerate([
        "El servicio será instalado en un plazo máximo de 5 días hábiles.",
        "El pago mensual es de Bs. 189 a realizarse los primeros 5 días de cada mes.",
        "Cualquier incumplimiento será notificado con 30 días de anticipación.",
    ],1):
        doc2.add_paragraph(f"{i}. {c}")
    path_contrato = f"{out_dir}/contrato_servicio.docx"
    doc2.save(path_contrato)
    print(f"  Word: {path_contrato}")

    # Excel - Reporte de Diagnóstico
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Diagnóstico Técnico"
    ws.column_dimensions["A"].width = 30
    ws.column_dimensions["B"].width = 40

    header_fill = PatternFill("solid", fgColor="6366F1")
    header_font = Font(bold=True, color="FFFFFF")
    ws["A1"] = "Reporte de Diagnóstico Técnico"
    ws.merge_cells("A1:B1")
    ws["A1"].font = Font(bold=True, size=14)
    ws["A1"].alignment = Alignment(horizontal="center")

    ws["A3"] = "Campo"; ws["B3"] = "Valor"
    ws["A3"].font = header_font; ws["A3"].fill = header_fill
    ws["B3"].font = header_font; ws["B3"].fill = header_fill

    rows = [
        ("Técnico asignado","Carlos López"),
        ("Fecha diagnóstico", datetime.now().strftime('%d/%m/%Y %H:%M')),
        ("Tipo de falla","Pérdida de señal ONT"),
        ("Equipo afectado","Router ZTE F660"),
        ("Número de serie","ZTE2024ABC123"),
        ("Resultado diagnóstico","Falla en splitter de distribución"),
        ("Solución aplicada","Reemplazo de splitter y reconfiguración"),
        ("Tiempo de atención (min)","45"),
        ("¿Requiere visita técnica?","No"),
        ("Estado final","Resuelto"),
    ]
    for i,(k,v) in enumerate(rows, 4):
        ws[f"A{i}"] = k
        ws[f"B{i}"] = v

    path_excel = f"{out_dir}/diagnostico_tecnico.xlsx"
    wb.save(path_excel)
    print(f"  Excel: {path_excel}")

    # Excel - Listado de trámites
    wb2 = openpyxl.Workbook()
    ws2 = wb2.active
    ws2.title = "Trámites"
    headers = ["Código","Título","Estado","Workflow","Responsable","Fecha"]
    for col, h_text in enumerate(headers, 1):
        cell = ws2.cell(1, col, h_text)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="6366F1")
    tramites_data = [
        ("TR-001","Falla internet edificio Norte","EN_PROGRESO","Atención de Incidencia","Juan Pérez","11/06/2026"),
        ("TR-002","Sin señal en oficina central","PENDIENTE","Atención de Incidencia","Juan Pérez","11/06/2026"),
        ("TR-003","Velocidad muy baja - piso 3","COMPLETADO","Atención de Incidencia","Carlos López","10/06/2026"),
    ]
    for row, data in enumerate(tramites_data, 2):
        for col, val in enumerate(data, 1):
            ws2.cell(row, col, val)
    for col in range(1,7):
        ws2.column_dimensions[openpyxl.utils.get_column_letter(col)].width = 25
    path_excel2 = f"{out_dir}/listado_tramites.xlsx"
    wb2.save(path_excel2)
    print(f"  Excel: {path_excel2}")

    return path_word, path_contrato, path_excel, path_excel2

# ── 1. Crear formularios en los nodos existentes ─────────────────────────────
def crear_formularios(token):
    print("\n[1] Creando formularios en nodos del workflow de incidencia...")

    forms = [
        {
            "nodoId": N_RECEP,
            "workflowId": WF_INCIDENCIA,
            "title": "Recepción de Incidencia",
            "fields": [
                {"name": "Nombre del cliente",     "type": "text",     "order": 1},
                {"name": "Teléfono",               "type": "text",     "order": 2},
                {"name": "Dirección",              "type": "text",     "order": 3},
                {"name": "Tipo de falla",          "type": "text",     "order": 4},
                {"name": "Descripción del problema","type": "TEXT",     "order": 5},
                {"name": "Adjuntar evidencia",     "type": "file",     "order": 6},
            ]
        },
        {
            "nodoId": N_DIAG,
            "workflowId": WF_INCIDENCIA,
            "title": "Diagnóstico Técnico",
            "fields": [
                {"name": "Tipo de falla técnica",     "type": "text",     "order": 1},
                {"name": "Equipo afectado",           "type": "text",     "order": 2},
                {"name": "Número de serie",           "type": "text",     "order": 3},
                {"name": "Resultado del diagnóstico", "type": "TEXT", "order": 4},
                {"name": "¿Requiere visita técnica?", "type": "text",     "order": 5},
                {"name": "Reporte técnico",           "type": "file",     "order": 6},
            ]
        },
        {
            "nodoId": N_CIERRE,
            "workflowId": WF_INCIDENCIA,
            "title": "Cierre de Incidencia",
            "fields": [
                {"name": "Solución aplicada",        "type": "TEXT", "order": 1},
                {"name": "Tiempo de resolución (min)","type": "text",    "order": 2},
                {"name": "Observaciones finales",    "type": "TEXT", "order": 3},
                {"name": "Satisfacción del cliente", "type": "text",     "order": 4},
            ]
        },
        {
            "nodoId": N_PROG,
            "workflowId": WF_INCIDENCIA,
            "title": "Programación de Visita Técnica",
            "fields": [
                {"name": "Fecha de visita",           "type": "text", "order": 1},
                {"name": "Hora de visita",            "type": "text", "order": 2},
                {"name": "Técnico asignado",          "type": "text", "order": 3},
                {"name": "Materiales requeridos",     "type": "TEXT", "order": 4},
                {"name": "Observaciones",             "type": "TEXT", "order": 5},
            ]
        },
    ]

    for f in forms:
        r = post(f"{BASE}/forms", token, f)
        status = "OK" if r else "ERROR"
        print(f"  {status}: Form '{f['title']}'")
    return True

# ── 2. Crear trámites de ejemplo ──────────────────────────────────────────────
WF_INCIDENCIA_NAME = "Atención de Incidencia de Internet"

def crear_tramites(token):
    print("\n[2] Creando trámites variados (PENDIENTE / EN_PROGRESO / COMPLETADO / RECHAZADO)...")

    # Generar documentos si no existen
    docs_dir = "/tmp/viva_docs"
    doc_reporte = f"{docs_dir}/reporte_incidencia.docx"
    doc_diag    = f"{docs_dir}/diagnostico_tecnico.xlsx"
    doc_contrato = f"{docs_dir}/contrato_servicio.docx"
    if not os.path.exists(doc_reporte):
        gen_docs(docs_dir)

    # Pre-cargar documentos al servidor (se usan en varios trámites)
    print("  Subiendo documentos de ejemplo...")
    file_reporte  = upload_file(token, doc_reporte,  WF_INCIDENCIA)
    file_diag     = upload_file(token, doc_diag,     WF_INCIDENCIA)
    file_contrato = upload_file(token, doc_contrato, WF_INCIDENCIA)
    if file_reporte:  print(f"    OK: {file_reporte.get('fileName')}")
    if file_diag:     print(f"    OK: {file_diag.get('fileName')}")
    if file_contrato: print(f"    OK: {file_contrato.get('fileName')}")

    def fd(nombre, tel, dir_, falla, desc, evidencia=None, reporte=None):
        d = {
            "Nombre del cliente": nombre, "Teléfono": tel,
            "Dirección": dir_, "Tipo de falla": falla,
            "Descripción del problema": desc,
        }
        if evidencia: d["Adjuntar evidencia"] = evidencia
        if reporte:   d["Reporte técnico"]    = reporte
        return d

    specs = [
        # ── PENDIENTE: recién creados, no han avanzado del nodo Inicio ──────
        {
            "title": "Nueva solicitud - Edificio Central",
            "description": "Cliente recién reportó falla, pendiente de recepción",
            "formData": fd("Luis Mamani","71100001","Av. Camacho 123","Sin señal","Sin internet desde esta mañana"),
            "autoTransitionIds": [],
            "_post": None,
        },
        {
            "title": "Falla red Wi-Fi - Sala Conferencias",
            "description": "Sala de reuniones sin conectividad",
            "formData": fd("Sofía Ríos","71100002","Calle Mercado 45","Wi-Fi caído","Todos los dispositivos sin conexión"),
            "autoTransitionIds": [],
            "_post": None,
        },
        # ── EN_PROGRESO en Recepción: tomaron el primer paso ─────────────────
        {
            "title": "Lentitud extrema - Piso 5",
            "description": "Velocidad insuficiente para trabajar",
            "formData": fd("Marco Apaza","71100003","Edificio Hansa Piso 5","Velocidad baja",
                           "Descarga máxima de 0.5 Mbps todo el día",
                           evidencia=file_contrato),
            "autoTransitionIds": [T_INICIO_TO_RECEP],
            "_post": None,
        },
        {
            "title": "Sin señal ONT - Zona Oeste",
            "description": "Equipo óptico sin sincronización",
            "formData": fd("Patricia Huanca","71100004","Zona Oeste Calle 3 #77","Sin señal ONT",
                           "Luz roja permanente en el ONT"),
            "autoTransitionIds": [T_INICIO_TO_RECEP],
            "_post": None,
        },
        # ── EN_PROGRESO en Diagnóstico Técnico (con reporte adjunto) ─────────
        {
            "title": "Intermitencia señal - Torre B",
            "description": "Cortes frecuentes cada media hora",
            "formData": fd("Jorge Condori","71100005","Torres Sofer Torre B Piso 8","Intermitencia",
                           "Servicio se corta exactamente a las horas",
                           evidencia=file_reporte, reporte=file_diag),
            "autoTransitionIds": [T_INICIO_TO_RECEP, T_RECEP_TO_DIAG],
            "_post": None,
        },
        {
            "title": "Caída total - Área Industrial",
            "description": "Empresa sin internet desde ayer",
            "formData": fd("Empresa XYZ SRL","71100006","Parque Industrial Norte #12","Corte total",
                           "Sin servicio, afecta operaciones críticas",
                           evidencia=file_reporte),
            "autoTransitionIds": [T_INICIO_TO_RECEP, T_RECEP_TO_DIAG],
            "_post": None,
        },
        # ── EN_PROGRESO en Cierre (decisión: aceptar, pendiente de cerrar) ───
        {
            "title": "Falla splitter - Residencial Sur",
            "description": "Splitter de distribución dañado",
            "formData": fd("Elena Vargas","71100007","Residencial Sur Calle 8","Sin señal",
                           "Diagnóstico indica splitter defectuoso",
                           evidencia=file_reporte, reporte=file_diag),
            "autoTransitionIds": [T_INICIO_TO_RECEP, T_RECEP_TO_DIAG, f"{T_DIAG_TO_DEC}>>{T_DEC_ACEPTAR}"],
            "_post": None,
        },
        # ── EN_PROGRESO en Programación de Visita ────────────────────────────
        {
            "title": "Cable dañado - Av. 6 de Agosto",
            "description": "Daño físico en la línea exterior",
            "formData": fd("Pedro Chávez","71100008","Av. 6 de Agosto #234","Falla física",
                           "Cable roto visible en el poste",
                           evidencia=file_reporte),
            "autoTransitionIds": [T_INICIO_TO_RECEP, T_RECEP_TO_DIAG, f"{T_DIAG_TO_DEC}>>{T_DEC_RECHAZAR}"],
            "_post": None,
        },
        # ── COMPLETADO: llegaron al nodo Fin (con documentos completos) ───────
        {
            "title": "Internet restaurado - Edificio Norte",
            "description": "Falla resuelta, servicio normalizado",
            "formData": fd("Roberto Mamani","71100009","Edificio Norte Piso 2","Sin señal",
                           "Resuelto: reemplazo de splitter",
                           evidencia=file_reporte, reporte=file_diag),
            "autoTransitionIds": [T_INICIO_TO_RECEP, T_RECEP_TO_DIAG,
                                  f"{T_DIAG_TO_DEC}>>{T_DEC_ACEPTAR}", T_CIERRE_TO_FIN],
            "_post": None,
        },
        {
            "title": "Velocidad normalizada - Oficina Central",
            "description": "Velocidad restaurada al 100%",
            "formData": fd("Ana Quispe","71100010","Calle Loayza 456","Velocidad baja",
                           "Resuelto: limpieza de conectores",
                           evidencia=file_reporte, reporte=file_diag),
            "autoTransitionIds": [T_INICIO_TO_RECEP, T_RECEP_TO_DIAG,
                                  f"{T_DIAG_TO_DEC}>>{T_DEC_ACEPTAR}", T_CIERRE_TO_FIN],
            "_post": None,
        },
        {
            "title": "Visita técnica completada - Zona Norte",
            "description": "Técnico realizó visita y resolvió el problema",
            "formData": fd("Carlos Flores","71100011","Zona Norte Av. Arce","Falla física",
                           "Cable exterior reemplazado",
                           evidencia=file_reporte, reporte=file_diag),
            "autoTransitionIds": [T_INICIO_TO_RECEP, T_RECEP_TO_DIAG,
                                  f"{T_DIAG_TO_DEC}>>{T_DEC_RECHAZAR}", T_PROG_TO_FIN],
            "_post": None,
        },
        # ── RECHAZADO: se submite y luego se llama /reject ───────────────────
        {
            "title": "Reporte falso - sin falla detectada",
            "description": "El técnico verificó y no hay ninguna falla",
            "formData": fd("Anónimo","71100012","Zona Central","Sin falla","El cliente reportó falla pero el equipo funciona bien"),
            "autoTransitionIds": [T_INICIO_TO_RECEP, T_RECEP_TO_DIAG],
            "_post": {"action": "reject", "reason": "Diagnóstico confirmó que no existe falla técnica. El equipo funciona correctamente."},
        },
        {
            "title": "Solicitud duplicada - ya existe TRM anterior",
            "description": "El cliente abrió dos tickets por la misma falla",
            "formData": fd("Luis Mamani","71100013","Av. Camacho 123","Duplicado","Trámite duplicado, ya existe uno activo"),
            "autoTransitionIds": [T_INICIO_TO_RECEP],
            "_post": {"action": "reject", "reason": "Solicitud duplicada. Ya existe un trámite activo para este cliente y dirección."},
        },
    ]

    created = []
    for t in specs:
        body = {k: v for k, v in t.items() if not k.startswith("_")}
        body["workflowId"] = WF_INCIDENCIA
        r = post(f"{BASE}/tramites/submit", token, body)
        if r:
            tid  = r.get("id", "?")
            code = r.get("code", "?")
            stat = r.get("status", "?")

            post_action = t.get("_post")
            if post_action and tid != "?":
                if post_action["action"] == "reject":
                    rj = requests.post(
                        f"{BASE}/activities/{tid}/reject",
                        json={"reason": post_action["reason"]},
                        headers=h(token)
                    )
                    stat = "RECHAZADO" if rj.ok else f"reject-err-{rj.status_code}"

            print(f"  [{code}] {stat:<12} '{t['title']}'")
            created.append(r)
        time.sleep(0.2)

    totales = {"PENDIENTE":0,"EN_PROGRESO":0,"COMPLETADO":0,"RECHAZADO":0}
    for c in created:
        s = c.get("status","?")
        if s in totales: totales[s] += 1
    print(f"  → PENDIENTE:{totales['PENDIENTE']} EN_PROGRESO:{totales['EN_PROGRESO']} "
          f"COMPLETADO:{totales['COMPLETADO']} RECHAZADO:{totales['RECHAZADO']}")
    return created

# ── 3. Crear nuevo workflow "Gestión de Contratos" ───────────────────────────
def crear_workflow_contratos(token):
    print("\n[3] Creando workflow 'Gestión de Contratos de Servicio'...")

    # Crear workflow
    wf = post(f"{BASE}/workflows", token, {
        "name": "Gestión de Contratos de Servicio",
        "description": "Proceso para la contratación de planes de internet: desde la solicitud "
                       "del cliente hasta la activación del servicio, incluyendo verificación "
                       "de cobertura, firma de contrato y habilitación técnica.",
        "companyId": COMPANY_ID,
    })
    if not wf: return None
    WF_ID = wf["id"]
    print(f"  Workflow creado: {WF_ID}")

    # Crear nodos (con posiciones para el diagrama visual)
    def nodo(name, ntype, dept=None, role=None, order=1, avg=30, x=0.0, y=0.0):
        return post(f"{BASE}/workflow-nodos", token, {
            "workflowId": WF_ID,
            "name": name, "nodeType": ntype, "order": order,
            "responsibleDepartmentId": dept, "responsibleJobRoleId": role,
            "avgMinutes": avg, "requiresForm": ntype == "proceso",
            "posX": x, "posY": y,
        })

    n_inicio  = nodo("Inicio",                    "inicio",   order=1,  x=100,  y=250)
    n_sol     = nodo("Solicitud de Contrato",      "proceso",  DEPT_ATENCION, ROLE_RECEP,  2, 30,  x=280,  y=250)
    n_ver     = nodo("Verificación de Cobertura",  "proceso",  DEPT_SOPORTE,  ROLE_TEC,    3, 60,  x=500,  y=250)
    n_dec     = nodo("¿Tiene cobertura?",          "decision", order=4,  x=720,  y=250)
    n_firma   = nodo("Firma de Contrato",          "proceso",  DEPT_ATENCION, ROLE_RECEP,  5, 20,  x=940,  y=130)
    n_activ   = nodo("Activación del Servicio",    "proceso",  DEPT_SOPORTE,  ROLE_TEC,    6, 90,  x=1160, y=130)
    n_fin_ok  = nodo("Servicio Activado",          "fin",      order=7,  x=1380, y=130)
    n_notif   = nodo("Notificación Sin Cobertura", "proceso",  DEPT_ATENCION, ROLE_RECEP,  8, 15,  x=940,  y=380)
    n_fin_no  = nodo("Sin Cobertura - Cerrado",    "fin",      order=9,  x=1160, y=380)

    nodes = {k:v for k,v in {
        "inicio":n_inicio,"solicitud":n_sol,"verificacion":n_ver,
        "decision":n_dec,"firma":n_firma,"activacion":n_activ,
        "fin_ok":n_fin_ok,"notif":n_notif,"fin_no":n_fin_no
    }.items() if v}
    print(f"  Nodos creados: {len(nodes)}/9")

    # Crear transiciones
    def trans(from_id, to_id, name=""):
        return post(f"{BASE}/workflow-transitions", token, {
            "workflowId": WF_ID,
            "fromNodoId": from_id, "toNodoId": to_id, "name": name
        })

    req = ["inicio","solicitud","verificacion","decision","firma","activacion","fin_ok","notif","fin_no"]
    if all(k in nodes for k in req):
        trans(nodes["inicio"]["id"],       nodes["solicitud"]["id"])
        trans(nodes["solicitud"]["id"],    nodes["verificacion"]["id"])
        trans(nodes["verificacion"]["id"], nodes["decision"]["id"])
        trans(nodes["decision"]["id"],     nodes["firma"]["id"],    "Tiene cobertura")
        trans(nodes["decision"]["id"],     nodes["notif"]["id"],    "Sin cobertura")
        trans(nodes["firma"]["id"],        nodes["activacion"]["id"])
        trans(nodes["activacion"]["id"],   nodes["fin_ok"]["id"])
        trans(nodes["notif"]["id"],        nodes["fin_no"]["id"])
        print("  Transiciones creadas: 8")

    # Crear formularios para cada nodo proceso
    if "solicitud" in nodes:
        post(f"{BASE}/forms", token, {
            "nodoId": nodes["solicitud"]["id"], "workflowId": WF_ID,
            "title": "Solicitud de Contrato",
            "fields": [
                {"name": "Nombre del cliente",    "type": "text",     "order": 1},
                {"name": "CI / RUT",              "type": "text",     "order": 2},
                {"name": "Teléfono de contacto",  "type": "text",     "order": 3},
                {"name": "Dirección de instalación","type": "text",   "order": 4},
                {"name": "Plan solicitado",        "type": "text",     "order": 5},
                {"name": "Fecha requerida",        "type": "text",     "order": 6},
                {"name": "Observaciones",          "type": "TEXT", "order": 7},
            ]
        })
    if "verificacion" in nodes:
        post(f"{BASE}/forms", token, {
            "nodoId": nodes["verificacion"]["id"], "workflowId": WF_ID,
            "title": "Verificación de Cobertura",
            "fields": [
                {"name": "¿Tiene cobertura?",          "type": "text",     "order": 1},
                {"name": "Tipo de infraestructura",    "type": "text",     "order": 2},
                {"name": "Velocidad máxima disponible","type": "text",     "order": 3},
                {"name": "Distancia a nodo más cercano","type": "text",    "order": 4},
                {"name": "Observaciones técnicas",     "type": "TEXT", "order": 5},
                {"name": "Informe de cobertura",       "type": "file",     "order": 6},
            ]
        })
    if "firma" in nodes:
        post(f"{BASE}/forms", token, {
            "nodoId": nodes["firma"]["id"], "workflowId": WF_ID,
            "title": "Firma de Contrato",
            "fields": [
                {"name": "Número de contrato",  "type": "text",  "order": 1},
                {"name": "Plan contratado",     "type": "text",  "order": 2},
                {"name": "Monto mensual (Bs.)", "type": "text",  "order": 3},
                {"name": "Fecha de firma",      "type": "text",  "order": 4},
                {"name": "Duración (meses)",    "type": "text",  "order": 5},
                {"name": "Contrato firmado",    "type": "file",  "order": 6},
            ]
        })
    if "activacion" in nodes:
        post(f"{BASE}/forms", token, {
            "nodoId": nodes["activacion"]["id"], "workflowId": WF_ID,
            "title": "Activación del Servicio",
            "fields": [
                {"name": "Fecha de activación",   "type": "text", "order": 1},
                {"name": "IP asignada",           "type": "text", "order": 2},
                {"name": "Equipo instalado",      "type": "text", "order": 3},
                {"name": "Número de serie",       "type": "text", "order": 4},
                {"name": "Velocidad medida (Mbps)","type": "text","order": 5},
                {"name": "Prueba de conexión OK", "type": "text", "order": 6},
            ]
        })
    # Formulario para nodo notificación
    if "notif" in nodes:
        post(f"{BASE}/forms", token, {
            "nodoId": nodes["notif"]["id"], "workflowId": WF_ID,
            "title": "Notificación Sin Cobertura",
            "fields": [
                {"name": "Motivo de rechazo",       "type": "TEXT", "order": 1},
                {"name": "Alternativas ofrecidas",  "type": "TEXT", "order": 2},
                {"name": "Fecha de notificación",   "type": "DATE", "order": 3},
            ]
        })
    print("  Formularios creados: 5")
    return WF_ID

# ── MAIN ──────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser()
    parser.add_argument("--skip-docs",     action="store_true")
    parser.add_argument("--skip-forms",    action="store_true")
    parser.add_argument("--skip-tramites", action="store_true")
    parser.add_argument("--skip-workflow", action="store_true")
    parser.add_argument("--delete-wf",    default=None, help="Delete orphaned workflow before creating")
    args = parser.parse_args()

    print("=" * 60)
    print("SEED VIVA - Workflow Manager")
    print("=" * 60)

    token = login("admin@viva.com", "12345")
    print(f"Login OK: admin@viva.com")

    if args.delete_wf:
        r = requests.delete(f"{BASE}/workflows/{args.delete_wf}", headers=h(token))
        print(f"  Deleted orphaned workflow {args.delete_wf}: {r.status_code}")

    if not args.skip_docs:
        print("\n[0] Generando documentos Word/Excel...")
        docs = gen_docs()
    else:
        print("\n[0] Documentos: omitido")
        docs = []

    if not args.skip_forms:
        crear_formularios(token)
    else:
        print("\n[1] Formularios: omitido")

    tramites = []
    if not args.skip_tramites:
        tramites = crear_tramites(token)
    else:
        print("\n[2] Trámites: omitido")

    wf_id = None
    if not args.skip_workflow:
        wf_id = crear_workflow_contratos(token)
    else:
        print("\n[3] Workflow contratos: omitido")

    print("\n" + "=" * 60)
    print("RESUMEN")
    print("=" * 60)
    print(f"  Documentos generados: 4 (en /tmp/viva_docs/)")
    print(f"  Formularios creados:  4 (workflow de incidencia)")
    print(f"  Trámites creados:     {len(tramites)}")
    print(f"  Workflow nuevo:       Gestión de Contratos ({wf_id})")
    print("\nListo! Recargá la app.")
