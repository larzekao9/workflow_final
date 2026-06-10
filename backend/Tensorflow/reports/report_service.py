"""
report_service.py
Genera reportes en Word o Excel a partir de una lista de filas.
"""

import io
import logging
from datetime import datetime

logger = logging.getLogger(__name__)

COLUMN_LABELS = {
    "tramiteId":      "ID Trámite",
    "code":           "Código",
    "title":          "Título",
    "workflowName":   "Workflow",
    "departmentName": "Departamento",
    "status":         "Estado",
    "userName":       "Usuario",
    "assignedToName": "Asignado a",
    "createdByName":  "Creado por",
    "createdAt":      "Fecha",
}


def _group_rows(rows: list[dict], group_by: str | None) -> list[tuple[str | None, list[dict]]]:
    """Agrupa filas por el campo indicado. Retorna [(group_label, [rows]), ...]."""
    if not group_by:
        return [(None, rows)]
    groups: dict[str, list[dict]] = {}
    for row in rows:
        key = str(row.get(group_by) or "Sin valor")
        groups.setdefault(key, []).append(row)
    return list(groups.items())


def generate_word(title: str, columns: list[str], rows: list[dict], group_by: str | None = None) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def _set_cell_bg(cell, hex_color: str):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tcPr.append(shd)

    doc = Document()

    h = doc.add_heading(title, level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h.runs[0].font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)

    sub = doc.add_paragraph(f"Generado el {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(9)
    sub.runs[0].font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
    doc.add_paragraph()

    if not rows:
        doc.add_paragraph("No se encontraron datos para los criterios seleccionados.")
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    labels = [COLUMN_LABELS.get(c, c) for c in columns]
    groups = _group_rows(rows, group_by)
    total  = 0

    for group_label, group_rows in groups:
        if group_label is not None:
            grp_p = doc.add_paragraph(f"{COLUMN_LABELS.get(group_by, group_by)}: {group_label} ({len(group_rows)})")
            grp_p.runs[0].bold = True
            grp_p.runs[0].font.size = Pt(10)
            grp_p.runs[0].font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)

        table = doc.add_table(rows=1 + len(group_rows), cols=len(columns))
        table.style = "Table Grid"

        for i, label in enumerate(labels):
            cell = table.rows[0].cells[i]
            cell.text = label
            run = cell.paragraphs[0].runs[0]
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            _set_cell_bg(cell, "4F46E5")

        for r_idx, row_data in enumerate(group_rows):
            row_cells = table.rows[r_idx + 1].cells
            for c_idx, col in enumerate(columns):
                val = str(row_data.get(col, "") or "")
                row_cells[c_idx].text = val
                runs = row_cells[c_idx].paragraphs[0].runs
                if runs:
                    runs[0].font.size = Pt(8)
                if r_idx % 2 == 0:
                    _set_cell_bg(row_cells[c_idx], "EEF2FF")

        total += len(group_rows)
        if group_label is not None:
            doc.add_paragraph()

    doc.add_paragraph()
    footer = doc.add_paragraph(f"Total: {total} registro(s)")
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generate_excel(title: str, columns: list[str], rows: list[dict], group_by: str | None = None) -> bytes:
    import openpyxl
    from openpyxl.styles import (Font, PatternFill, Alignment,
                                  Border, Side)

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Reporte"

    # Título fusionado
    ws.merge_cells(start_row=1, start_column=1,
                   end_row=1, end_column=max(len(columns), 1))
    title_cell = ws.cell(row=1, column=1, value=title)
    title_cell.font      = Font(bold=True, size=14, color="4F46E5")
    title_cell.alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[1].height = 28

    # Subtítulo
    ws.merge_cells(start_row=2, start_column=1,
                   end_row=2, end_column=max(len(columns), 1))
    sub_cell = ws.cell(row=2, column=1,
                       value=f"Generado: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    sub_cell.font      = Font(size=9, color="94A3B8")
    sub_cell.alignment = Alignment(horizontal="center")

    # Encabezados (fila 3)
    HDR_FILL  = PatternFill("solid", fgColor="4F46E5")
    HDR_FONT  = Font(bold=True, color="FFFFFF", size=10)
    HDR_ALIGN = Alignment(horizontal="center", vertical="center", wrap_text=True)
    thin      = Side(style="thin", color="CBD5E1")
    border    = Border(left=thin, right=thin, top=thin, bottom=thin)

    labels = [COLUMN_LABELS.get(c, c) for c in columns]
    for ci, label in enumerate(labels, start=1):
        cell             = ws.cell(row=3, column=ci, value=label)
        cell.fill        = HDR_FILL
        cell.font        = HDR_FONT
        cell.alignment   = HDR_ALIGN
        cell.border      = border
        ws.column_dimensions[cell.column_letter].width = 20

    ws.row_dimensions[3].height = 22

    # Filas de datos (con soporte de groupBy)
    ALT_FILL   = PatternFill("solid", fgColor="EEF2FF")
    GRP_FILL   = PatternFill("solid", fgColor="E0E7FF")
    DATA_FONT  = Font(size=9)
    DATA_ALIGN = Alignment(vertical="center")
    GRP_FONT   = Font(size=10, bold=True, color="4F46E5")

    groups    = _group_rows(rows, group_by)
    cur_row   = 4
    total     = 0

    for group_label, group_rows in groups:
        if group_label is not None:
            # Fila de encabezado de grupo
            grp_label_col = COLUMN_LABELS.get(group_by, group_by) if group_by else ""
            merge_end = max(len(columns), 1)
            ws.merge_cells(start_row=cur_row, start_column=1, end_row=cur_row, end_column=merge_end)
            gc = ws.cell(row=cur_row, column=1,
                         value=f"{grp_label_col}: {group_label}  ({len(group_rows)} registro(s))")
            gc.font      = GRP_FONT
            gc.fill      = GRP_FILL
            gc.alignment = Alignment(vertical="center")
            ws.row_dimensions[cur_row].height = 18
            cur_row += 1

        for ri, row_data in enumerate(group_rows):
            fill = ALT_FILL if ri % 2 == 0 else None
            for ci, col in enumerate(columns, start=1):
                val  = row_data.get(col, "")
                cell = ws.cell(row=cur_row, column=ci, value=str(val) if val else "")
                cell.font      = DATA_FONT
                cell.alignment = DATA_ALIGN
                cell.border    = border
                if fill:
                    cell.fill = fill
            ws.row_dimensions[cur_row].height = 16
            cur_row += 1
            total  += 1

    ws.cell(row=cur_row, column=1, value=f"Total: {total} registro(s)").font = Font(
        size=9, color="64748B", bold=True
    )

    # Freeze encabezados
    ws.freeze_panes = "A4"

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()
