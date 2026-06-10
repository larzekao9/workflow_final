"""
document_reader.py
Extrae texto plano de archivos Word (.docx), PDF, y TXT.
"""

import io
import logging

logger = logging.getLogger(__name__)


def extract_text(filename: str, content: bytes) -> str:
    """
    Dado el nombre del archivo y su contenido en bytes,
    devuelve el texto extraído como string.
    """
    ext = filename.rsplit(".", 1)[-1].lower()

    try:
        if ext == "docx":
            return _read_docx(content)
        elif ext == "pdf":
            return _read_pdf(content)
        elif ext in ("txt", "text"):
            return content.decode("utf-8", errors="ignore")
        else:
            logger.warning(f"Tipo de archivo no soportado: {ext}")
            return ""
    except Exception as e:
        logger.error(f"Error leyendo {filename}: {e}")
        return ""


def _read_docx(content: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(content))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # También extraer texto de tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())
    return "\n".join(paragraphs)


def _read_pdf(content: bytes) -> str:
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            pages = []
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    pages.append(text)
            return "\n".join(pages)
    except Exception:
        # Fallback a PyPDF2
        import PyPDF2
        reader = PyPDF2.PdfReader(io.BytesIO(content))
        return "\n".join(
            page.extract_text() or "" for page in reader.pages
        )
