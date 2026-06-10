import io
import json
from datetime import datetime
from typing import Any

from fastapi import HTTPException
from fastapi.responses import StreamingResponse

from ai_common import DIAGRAM_MODEL, HAIKU_MODEL, call_claude, get_db, parse_json_response


REPORT_SYSTEM_PROMPT = """
Eres un agente experto en análisis de datos y generación de reportes sobre una base de datos MongoDB.
Interpreta solicitudes en lenguaje natural (pueden venir de voz convertida a texto) y genera un
pipeline de agregación de MongoDB para obtener los datos pedidos.

Responde SIEMPRE con JSON puro, sin markdown, sin texto fuera del JSON.

Formato obligatorio:
{
  "collection": "nombre_de_la_coleccion",
  "pipeline": [...],
  "format": "screen|pdf|excel|word",
  "title": "Titulo del reporte",
  "columns": ["campo1", "campo2"],
  "warnings": []
}

Reglas de formato de salida (infiere de la solicitud):
- "pantalla", "ver", "mostrar", "listar" -> "screen"
- "pdf" -> "pdf"
- "excel", "hoja de calculo", "xlsx", "planilla" -> "excel"
- "word", "documento", "doc" -> "word"
- Si no especifica formato, usa "screen"

Reglas de pipeline:
- Usa SOLO las colecciones listadas en el contexto.
- Si hay companyId, agrega {"$match": {"companyId": "<companyId>"}} como primera etapa.
- Soporta: $group, $match, $sort, $limit, $project, $unwind, $lookup, $count.
- Promedio: {"$avg": "$campo"} dentro de $group.
- Suma: {"$sum": "$campo"} dentro de $group.
- Conteo: {"$sum": 1} dentro de $group, o usa $count.
- Agrupación: {"$group": {"_id": "$campo", "total": {"$sum": 1}}}.
- Ordenar ASC: {"$sort": {"campo": 1}}, DESC: {"$sort": {"campo": -1}}.
- Siempre agrega $project al final para seleccionar solo los campos relevantes y excluir _id si no es necesario.
- Para fechas usa $dateToString, $year, $month, $dayOfMonth.
- Si el usuario pide "los últimos N", ordena desc por fecha y limita a N.
- columns debe listar los campos finales que aparecerán en el resultado (después del $project).
- Si la solicitud es ambigua, genera un pipeline simple y agrega warning.
- No inventes colecciones o campos que no existan en el esquema.

Ejemplos de solicitudes y pipelines:
- "muéstrame todos los trámites" -> $match companyId, $project campos relevantes
- "promedio de duración agrupado por estado" -> $group por estado, $avg de duración
- "los últimos 10 usuarios registrados" -> $sort createdAt desc, $limit 10
- "total de trámites por flujo de trabajo en excel" -> $group por workflowId, $sum, format excel
- "cuántos documentos se subieron por departamento" -> $group por departmentId/departmentName, $sum

Colecciones disponibles con sus campos (esquema inferido de datos reales):
"""


def get_collection_schemas(db: Any, company_id: str | None) -> str:
    try:
        collection_names = db.list_collection_names()
    except Exception:
        return "(no se pudo leer el esquema)"

    skip = {"fs.files", "fs.chunks", "system.indexes"}
    lines: list[str] = []

    for name in collection_names:
        if name in skip:
            continue
        try:
            sample_filter = {"companyId": company_id} if company_id else {}
            anchor = db[name].find_one(sample_filter, {"_id": 1}) or db[name].find_one({}, {"_id": 1})
            if anchor is None:
                continue
            full = db[name].find_one({"_id": anchor["_id"]})
            if full is None:
                continue
            fields = _extract_field_types(full)
            lines.append(f"- '{name}': {json.dumps(fields, ensure_ascii=False)}")
        except Exception:
            pass

    return "\n".join(lines) if lines else "(sin datos de esquema)"


def _extract_field_types(doc: dict, depth: int = 0) -> dict[str, str]:
    if depth > 2:
        return {}
    result: dict[str, str] = {}
    for key, value in doc.items():
        if key.startswith("__"):
            continue
        if isinstance(value, dict) and depth < 2:
            nested = _extract_field_types(value, depth + 1)
            for sub_key, sub_type in nested.items():
                result[f"{key}.{sub_key}"] = sub_type
        elif isinstance(value, list):
            result[key] = "array"
        elif isinstance(value, bool):
            result[key] = "boolean"
        elif isinstance(value, (int, float)):
            result[key] = "number"
        elif isinstance(value, str):
            result[key] = "string"
        elif value is None:
            result[key] = "null"
        else:
            result[key] = type(value).__name__
    return result


def process_report_request(body: dict[str, Any]) -> Any:
    transcript = str(body.get("transcript") or body.get("text") or "").strip()
    company_id = str(body.get("companyId") or "").strip() or None

    if not transcript:
        raise HTTPException(status_code=400, detail="Debes enviar el texto o transcripción de la solicitud")

    db = get_db()
    schema_context = get_collection_schemas(db, company_id)
    system_prompt = REPORT_SYSTEM_PROMPT + "\n" + schema_context

    user_content = f"Solicitud: {transcript}"
    if company_id:
        user_content += f"\ncompanyId del usuario: {company_id}"

    messages = [{"role": "user", "content": user_content}]

    try:
        raw = call_claude(system_prompt, HAIKU_MODEL, 4096, messages)
    except HTTPException as exc:
        if exc.status_code != 402:
            raise
        raw = call_claude(system_prompt, DIAGRAM_MODEL, 4096, messages)

    plan = parse_json_response(raw)

    if not plan or not plan.get("collection") or not isinstance(plan.get("pipeline"), list):
        raise HTTPException(status_code=422, detail="No se pudo interpretar la solicitud como un reporte válido")

    collection_name = str(plan["collection"]).strip()
    pipeline: list = plan["pipeline"]
    report_format = str(plan.get("format") or "screen").lower().strip()
    title = str(plan.get("title") or "Reporte").strip()
    columns: list[str] = [str(c) for c in (plan.get("columns") or [])]
    warnings: list[str] = [str(w) for w in (plan.get("warnings") or []) if str(w).strip()]

    available = db.list_collection_names()
    if collection_name not in available:
        raise HTTPException(status_code=422, detail=f"Colección '{collection_name}' no existe. Disponibles: {available}")

    try:
        cursor = db[collection_name].aggregate(pipeline, allowDiskUse=True)
        results: list[dict[str, Any]] = []
        for doc in cursor:
            doc.pop("_id", None)
            results.append(json.loads(json.dumps(doc, default=str)))
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Error ejecutando consulta MongoDB: {exc}")

    if not columns and results:
        columns = list(results[0].keys())

    if report_format == "excel":
        return _generate_excel(title, columns, results)
    if report_format == "pdf":
        return _generate_pdf(title, columns, results)
    if report_format == "word":
        return _generate_word(title, columns, results)

    return {
        "title": title,
        "columns": columns,
        "data": results,
        "count": len(results),
        "format": "screen",
        "warnings": warnings,
    }


# ── Excel ────────────────────────────────────────────────────────────────────

def _generate_excel(title: str, columns: list[str], data: list[dict]) -> StreamingResponse:
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Font, PatternFill

    wb = Workbook()
    ws = wb.active
    ws.title = title[:31]

    num_cols = max(len(columns), 1)
    last_col_letter = _col_letter(num_cols)

    # Title row
    ws.merge_cells(f"A1:{last_col_letter}1")
    cell = ws["A1"]
    cell.value = title
    cell.font = Font(bold=True, size=14, color="1F4E79")
    cell.alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[1].height = 22

    # Date row
    ws.merge_cells(f"A2:{last_col_letter}2")
    date_cell = ws["A2"]
    date_cell.value = f"Generado: {datetime.now().strftime('%d/%m/%Y %H:%M')}"
    date_cell.font = Font(italic=True, size=9, color="666666")
    date_cell.alignment = Alignment(horizontal="right")

    # Header row
    header_font = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill(start_color="1F4E79", end_color="1F4E79", fill_type="solid")
    for col_idx, col_name in enumerate(columns, 1):
        c = ws.cell(row=3, column=col_idx, value=col_name)
        c.font = header_font
        c.fill = header_fill
        c.alignment = Alignment(horizontal="center")

    # Data rows
    alt_fill = PatternFill(start_color="EBF3FB", end_color="EBF3FB", fill_type="solid")
    for row_idx, row in enumerate(data, 4):
        fill = alt_fill if row_idx % 2 == 0 else None
        for col_idx, col_name in enumerate(columns, 1):
            c = ws.cell(row=row_idx, column=col_idx, value=row.get(col_name, ""))
            if fill:
                c.fill = fill

    # Auto-width
    for col in ws.columns:
        max_len = 0
        for c in col:
            try:
                if c.value:
                    max_len = max(max_len, len(str(c.value)))
            except Exception:
                pass
        ws.column_dimensions[col[0].column_letter].width = min(max_len + 3, 50)

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    filename = f"{_safe_name(title)}_{_ts()}.xlsx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


def _col_letter(n: int) -> str:
    letter = ""
    while n > 0:
        n, remainder = divmod(n - 1, 26)
        letter = chr(65 + remainder) + letter
    return letter


# ── PDF ──────────────────────────────────────────────────────────────────────

def _generate_pdf(title: str, columns: list[str], data: list[dict]) -> StreamingResponse:
    try:
        from fpdf import FPDF
    except ImportError:
        raise HTTPException(status_code=501, detail="PDF no disponible: instala fpdf2 en el servidor")

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    # Title
    pdf.set_font("Helvetica", "B", 16)
    pdf.set_text_color(31, 78, 121)
    pdf.cell(0, 10, title, new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.ln(2)

    # Date
    pdf.set_font("Helvetica", "I", 9)
    pdf.set_text_color(100, 100, 100)
    pdf.cell(0, 6, f"Generado: {datetime.now().strftime('%d/%m/%Y %H:%M')}", new_x="LMARGIN", new_y="NEXT", align="R")
    pdf.ln(4)

    if not columns:
        pdf.set_font("Helvetica", "", 11)
        pdf.set_text_color(0, 0, 0)
        pdf.cell(0, 8, "Sin datos.", new_x="LMARGIN", new_y="NEXT")
    else:
        usable_width = pdf.w - pdf.l_margin - pdf.r_margin
        col_w = min(usable_width / len(columns), 55)

        # Header
        pdf.set_fill_color(31, 78, 121)
        pdf.set_text_color(255, 255, 255)
        pdf.set_font("Helvetica", "B", 9)
        for col in columns:
            pdf.cell(col_w, 8, str(col)[:22], border=1, fill=True, align="C")
        pdf.ln()

        # Rows
        pdf.set_font("Helvetica", "", 8)
        for i, row in enumerate(data):
            if i % 2 == 0:
                pdf.set_fill_color(235, 243, 251)
                fill = True
            else:
                pdf.set_fill_color(255, 255, 255)
                fill = True
            pdf.set_text_color(0, 0, 0)
            for col in columns:
                pdf.cell(col_w, 7, str(row.get(col, ""))[:28], border=1, fill=fill)
            pdf.ln()

    buf = io.BytesIO(pdf.output())
    buf.seek(0)
    filename = f"{_safe_name(title)}_{_ts()}.pdf"
    return StreamingResponse(
        buf,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ── Word ─────────────────────────────────────────────────────────────────────

def _generate_word(title: str, columns: list[str], data: list[dict]) -> StreamingResponse:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    doc = Document()

    heading = doc.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    date_para = doc.add_paragraph(f"Generado: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_para.runs[0].font.size = Pt(9)
    date_para.runs[0].font.color.rgb = RGBColor(100, 100, 100)
    doc.add_paragraph()

    if not columns:
        doc.add_paragraph("Sin datos para mostrar.")
    else:
        table = doc.add_table(rows=1, cols=len(columns))
        table.style = "Table Grid"

        # Header cells
        hdr_cells = table.rows[0].cells
        for idx, col_name in enumerate(columns):
            cell = hdr_cells[idx]
            cell.text = str(col_name)
            run = cell.paragraphs[0].runs[0]
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            _set_cell_bg(cell, "1F4E79")

        # Data rows
        for i, row in enumerate(data):
            row_cells = table.add_row().cells
            bg = "EBF3FB" if i % 2 == 0 else "FFFFFF"
            for idx, col_name in enumerate(columns):
                row_cells[idx].text = str(row.get(col_name, ""))
                _set_cell_bg(row_cells[idx], bg)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    filename = f"{_safe_name(title)}_{_ts()}.docx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


def _set_cell_bg(cell: Any, hex_color: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


# ── Helpers ───────────────────────────────────────────────────────────────────

def _safe_name(text: str) -> str:
    name = "".join(c for c in text if c.isalnum() or c in " _-").strip()
    return name or "reporte"


def _ts() -> str:
    return datetime.now().strftime("%Y%m%d_%H%M%S")
